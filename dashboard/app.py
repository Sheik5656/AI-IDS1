import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime
import time
import sys
import os
import threading
import tempfile
import base64
import csv
import json
from io import StringIO, BytesIO
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add parent directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.packet_sniffer import PacketSniffer, FileAnalyzer
from src.deepseek_analyzer import DeepSeekAnalyzer

# Page config
st.set_page_config(
    page_title="AI Intrusion Detection System",
    page_icon="🛡️",
    layout="wide"
)

# ===== YOUR DEEPSEEK API KEY - HARDCODED =====
DEEPSEEK_API_KEY = "sk-4fc8743996e04434baa71ecd1cd973c3"
# =============================================

# Initialize session state
if 'sniffer' not in st.session_state:
    st.session_state.sniffer = PacketSniffer()
    st.session_state.sniffer.load_model()
    
if 'file_analyzer' not in st.session_state:
    st.session_state.file_analyzer = FileAnalyzer()
    
if 'analyzer' not in st.session_state:
    st.session_state.analyzer = DeepSeekAnalyzer(api_key=DEEPSEEK_API_KEY)
    
if 'detection_active' not in st.session_state:
    st.session_state.detection_active = False
    
if 'thread' not in st.session_state:
    st.session_state.thread = None
    
if 'file_analysis_results' not in st.session_state:
    st.session_state.file_analysis_results = []
    
if 'last_analysis' not in st.session_state:
    st.session_state.last_analysis = None

# Custom CSS
st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #1e3c2c, #0a1f1a);
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .alert-box {
        padding: 1rem;
        border-radius: 5px;
        margin: 0.5rem 0;
        border-left: 4px solid;
    }
    .critical { border-color: #ff4b4b; background-color: #2a1a1a; }
    .high { border-color: #ffa64b; background-color: #2a241a; }
    .medium { border-color: #ffd24b; background-color: #2a2a1a; }
    .low { border-color: #4bff4b; background-color: #1a2a1a; }
    .metric-card {
        background: linear-gradient(135deg, #1a2a2a, #0f1f1a);
        padding: 1.5rem;
        border-radius: 10px;
        text-align: center;
        border: 1px solid #2a4a3a;
    }
    .stButton button {
        width: 100%;
        background: #1a3a2a;
        color: #9eff9e;
        border: 1px solid #2a4a3a;
    }
    .stButton button:hover {
        background: #2a4a3a;
        border-color: #f0b823;
    }
    .file-upload-area {
        border: 2px dashed #2a4a3a;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        background: #0f1f1a;
        margin: 1rem 0;
    }
    .risk-meter {
        width: 100%;
        height: 20px;
        background: linear-gradient(90deg, #00ff00, #ffff00, #ff0000);
        border-radius: 10px;
        margin: 10px 0;
    }
    .hash-box {
        font-family: monospace;
        background: #0a1a1a;
        padding: 5px;
        border-radius: 3px;
        border: 1px solid #2a4a3a;
    }
    .download-btn {
        background-color: #1a3a2a;
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 5px;
        text-decoration: none;
        border: 1px solid #2a4a3a;
        display: inline-block;
        margin: 0.5rem;
        text-align: center;
    }
    .download-btn:hover {
        background-color: #2a4a3a;
    }
    .download-container {
        display: flex;
        gap: 1rem;
        justify-content: center;
        margin: 1rem 0;
        flex-wrap: wrap;
    }
    .section-header {
        color: #f0b823;
        margin-top: 2rem;
        margin-bottom: 1rem;
        border-bottom: 1px solid #2a4a3a;
        padding-bottom: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown('<div class="main-header">', unsafe_allow_html=True)
st.title("🛡️ AI-Powered Intrusion Detection System")
st.markdown('</div>', unsafe_allow_html=True)

# Create tabs
tab1, tab2, tab3 = st.tabs(["📡 Live Network Detection", "📁 File Upload Analysis", "📊 Reports"])

# ===== TAB 1: Live Network Detection =====
with tab1:
    # Sidebar
    with st.sidebar:
        st.image("https://img.icons8.com/fluency/96/artificial-intelligence.png", width=80)
        st.title("Control Panel")
        
        st.markdown("### 🎮 Detection Controls")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("▶️ START", use_container_width=True):
                if not st.session_state.detection_active:
                    st.session_state.detection_active = True
                    thread = threading.Thread(target=st.session_state.sniffer.start_sniffing)
                    thread.daemon = True
                    thread.start()
                    st.success("✅ Detection started!")
                    
        with col2:
            if st.button("⏹️ STOP", use_container_width=True):
                st.session_state.sniffer.stop_sniffing()
                st.session_state.detection_active = False
                st.warning("🛑 Detection stopped")
        
        st.divider()
        
        st.markdown("### 📊 System Status")
        if st.session_state.detection_active:
            st.markdown("🟢 **Active**")
        else:
            st.markdown("🔴 **Inactive**")
        
        st.markdown(f"📡 **Interface:** {st.session_state.sniffer.interface}")
        st.markdown(f"📦 **Mode:** {'Simulation' if st.session_state.sniffer.simulation_mode else 'Live'}")

    # Main dashboard
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        stats = st.session_state.sniffer.get_stats()
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Packets Captured", stats['packet_count'])
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Active Alerts", stats['alert_count'])
        st.markdown('</div>', unsafe_allow_html=True)

    with col3:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Packet Rate", f"{stats['packet_rate']}/s")
        st.markdown('</div>', unsafe_allow_html=True)

    with col4:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Unique IPs", stats['unique_ips'])
        st.markdown('</div>', unsafe_allow_html=True)

    # Charts
    st.markdown("### 📈 Live Traffic Analysis")
    col1, col2 = st.columns(2)

    with col1:
        # Protocol distribution
        protocols = {'TCP': np.random.randint(10, 50), 
                     'UDP': np.random.randint(5, 30),
                     'ICMP': np.random.randint(0, 10),
                     'Other': np.random.randint(0, 5)}
        
        fig = go.Figure(data=[go.Pie(
            labels=list(protocols.keys()),
            values=list(protocols.values()),
            hole=0.4,
            marker_colors=['#00ff00', '#ffaa00', '#ff4b4b', '#888888']
        )])
        fig.update_layout(
            title="Protocol Distribution",
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            font_color='#9eff9e'
        )
        st.plotly_chart(fig, use_container_width=True)

    with col2:
        # Packet rate over time
        rates = np.random.randint(5, 30, 20)
        fig = go.Figure(data=[go.Scatter(
            y=rates,
            mode='lines+markers',
            line=dict(color='#00ff00', width=2),
            marker=dict(color='#f0b823', size=6)
        )])
        fig.update_layout(
            title="Packet Rate (last 20 samples)",
            xaxis_title="Time",
            yaxis_title="Packets/sec",
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            font_color='#9eff9e'
        )
        st.plotly_chart(fig, use_container_width=True)

    # Alerts section
    st.markdown("### 🚨 Recent Alerts")

    alerts = st.session_state.sniffer.get_alerts()

    if alerts:
        for i, alert in enumerate(reversed(alerts[-10:])):
            confidence = alert.get('confidence', 0.5)
            
            if confidence > 0.8:
                level = "critical"
                level_text = "CRITICAL"
            elif confidence > 0.6:
                level = "high"
                level_text = "HIGH"
            else:
                level = "medium"
                level_text = "MEDIUM"
            
            attack_type = alert.get('attack_type', 'Intrusion Detection')
            
            col1, col2, col3 = st.columns([3, 2, 1])
            
            with col1:
                st.markdown(f"**{alert['timestamp']}**")
                st.markdown(f"📍 {alert['src_ip']} → {alert['dst_ip']}")
                st.markdown(f"**Type:** {attack_type}")
            
            with col2:
                st.markdown(f"📡 Protocol: {alert['protocol']}")
                st.markdown(f"🎯 Confidence: {alert['confidence']:.1%}")
            
            with col3:
                st.markdown(f"**{level_text}**")
            
            st.divider()
    else:
        st.info("No alerts detected yet. Start the detection system to begin monitoring.")

# ===== TAB 2: File Upload Analysis =====
with tab2:
    st.markdown("## 📁 Malicious File Detection")
    
    # File uploader
    with st.container():
        st.markdown('<div class="file-upload-area">', unsafe_allow_html=True)
        uploaded_file = st.file_uploader(
            "Choose a file to analyze", 
            type=['js', 'php', 'py', 'ps1', 'sql', 'log', 'txt', 'csv', 'exe', 'dll', 'pdf', 'doc', 'zip', 'pcap'],
            key="file_uploader",
            help="Upload any file type for analysis"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    if uploaded_file is not None:
        # Check if this file was already analyzed
        file_already_analyzed = False
        for result in st.session_state.file_analysis_results:
            if result['filename'] == uploaded_file.name and result['file_size'] == uploaded_file.size:
                file_already_analyzed = True
                analysis_result = result
                break
        
        if not file_already_analyzed:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name
            
            # Show analysis progress
            with st.spinner("🔍 Analyzing file for malicious content..."):
                # Analyze the file
                analysis_result = st.session_state.file_analyzer.analyze_file(tmp_path, uploaded_file.name)
                st.session_state.file_analysis_results.append(analysis_result)
                
                # Get AI analysis
                deepseek_input = {
                    'timestamp': analysis_result['timestamp'],
                    'src_ip': 'FILE_UPLOAD',
                    'dst_ip': 'LOCAL_SYSTEM',
                    'protocol': 'FILE',
                    'confidence': analysis_result['risk_score'] / 100,
                    'features': analysis_result
                }
                ai_result = st.session_state.analyzer.analyze_alert(deepseek_input)
                analysis_result['ai_analysis'] = ai_result
                st.session_state['last_analysis'] = ai_result
            
            # Clean up temp file
            os.unlink(tmp_path)
        else:
            st.info(f"✅ File '{uploaded_file.name}' already analyzed. Showing previous results.")
        
        # Display results
        col1, col2, col3 = st.columns(3)
        
        with col1:
            risk_color = "red" if analysis_result['risk_level'] in ['CRITICAL', 'HIGH'] else "orange" if analysis_result['risk_level'] == 'MEDIUM' else "green"
            st.markdown(f"### Risk Level: :{risk_color}[**{analysis_result['risk_level']}**]")
            st.markdown(f"**Risk Score:** {analysis_result['risk_score']}/100")
            
            # Risk meter
            st.markdown(f"""
            <div style="width:100%; height:20px; background:#ddd; border-radius:10px;">
                <div style="width:{analysis_result['risk_score']}%; height:20px; 
                     background:{'#ff0000' if analysis_result['risk_score']>70 else '#ffaa00' if analysis_result['risk_score']>30 else '#00ff00'}; 
                     border-radius:10px;"></div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("### File Details")
            st.markdown(f"**Filename:** {analysis_result['filename']}")
            st.markdown(f"**Size:** {analysis_result['file_size']:,} bytes")
            st.markdown(f"**Extension:** {analysis_result['file_type'].get('extension', 'Unknown')}")
            st.markdown(f"**Entropy:** {analysis_result['entropy']:.2f}")
        
        with col3:
            st.markdown("### File Hashes")
            st.markdown(f"**MD5:** `{analysis_result['hash']['md5'][:16]}...`")
            st.markdown(f"**SHA1:** `{analysis_result['hash']['sha1'][:16]}...`")
            st.markdown(f"**SHA256:** `{analysis_result['hash']['sha256'][:16]}...`")
        
        # Download single file report
        st.markdown("### 📥 Download This Report")
        col1, col2, col3, col4 = st.columns(4)
        
        # Prepare single file report data
        single_report = {
            'filename': analysis_result['filename'],
            'timestamp': analysis_result['timestamp'],
            'risk_level': analysis_result['risk_level'],
            'risk_score': analysis_result['risk_score'],
            'file_size': analysis_result['file_size'],
            'entropy': analysis_result['entropy'],
            'md5': analysis_result['hash']['md5'],
            'sha1': analysis_result['hash']['sha1'],
            'sha256': analysis_result['hash']['sha256'],
            'suspicious_patterns': analysis_result.get('suspicious_patterns', []),
            'attack_summary': analysis_result.get('attack_summary', 'None')
        }
        
        # Add AI analysis if available
        if 'ai_analysis' in analysis_result:
            single_report['ai_analysis'] = analysis_result['ai_analysis']
        
        # Convert to DataFrame for CSV
        single_df = pd.DataFrame([{
            'Filename': analysis_result['filename'],
            'Timestamp': analysis_result['timestamp'],
            'Risk Level': analysis_result['risk_level'],
            'Risk Score': analysis_result['risk_score'],
            'File Size': analysis_result['file_size'],
            'Entropy': f"{analysis_result['entropy']:.2f}",
            'MD5': analysis_result['hash']['md5'],
            'SHA256': analysis_result['hash']['sha256'][:16] + '...',
            'Patterns': len(analysis_result.get('suspicious_patterns', []))
        }])
        
        with col1:
            # Single CSV Download
            csv_data = single_df.to_csv(index=False)
            b64_csv = base64.b64encode(csv_data.encode()).decode()
            href_csv = f'<a href="data:file/csv;base64,{b64_csv}" download="{analysis_result["filename"]}_report.csv" class="download-btn">📥 CSV</a>'
            st.markdown(href_csv, unsafe_allow_html=True)
        
        with col2:
            # Single JSON Download
            json_data = json.dumps(single_report, indent=2, default=str)
            b64_json = base64.b64encode(json_data.encode()).decode()
            href_json = f'<a href="data:file/json;base64,{b64_json}" download="{analysis_result["filename"]}_report.json" class="download-btn">📥 JSON</a>'
            st.markdown(href_json, unsafe_allow_html=True)
        
        with col3:
            # Single TXT Download
            txt_data = f"""
===========================================
AI INTRUSION DETECTION SYSTEM - FILE REPORT
===========================================
Filename: {analysis_result['filename']}
Timestamp: {analysis_result['timestamp']}
Risk Level: {analysis_result['risk_level']}
Risk Score: {analysis_result['risk_score']}/100
File Size: {analysis_result['file_size']} bytes
Entropy: {analysis_result['entropy']:.2f}

FILE HASHES:
- MD5: {analysis_result['hash']['md5']}
- SHA1: {analysis_result['hash']['sha1']}
- SHA256: {analysis_result['hash']['sha256']}

SUSPICIOUS PATTERNS FOUND: {len(analysis_result.get('suspicious_patterns', []))}
"""
            if analysis_result.get('suspicious_patterns'):
                txt_data += "\nPATTERNS:\n"
                for p in analysis_result['suspicious_patterns']:
                    txt_data += f"- {p['type']}: {p['pattern']}\n"
            
            b64_txt = base64.b64encode(txt_data.encode()).decode()
            href_txt = f'<a href="data:file/txt;base64,{b64_txt}" download="{analysis_result["filename"]}_report.txt" class="download-btn">📥 TXT</a>'
            st.markdown(href_txt, unsafe_allow_html=True)
        
        with col4:
            # Single HTML Download
            html_data = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>File Analysis Report - {analysis_result['filename']}</title>
                <style>
                    body {{ font-family: Arial; background: #0a0c0f; color: #dbecdb; padding: 20px; }}
                    .container {{ max-width: 800px; margin: auto; }}
                    .header {{ background: #1e3c2c; padding: 20px; border-radius: 10px; text-align: center; }}
                    .risk-{analysis_result['risk_level'].lower()} {{ color: {'#ff4b4b' if analysis_result['risk_level'] in ['CRITICAL', 'HIGH'] else '#ffaa00' if analysis_result['risk_level'] == 'MEDIUM' else '#00ff00'}; }}
                    .section {{ background: #151e2a; padding: 15px; margin: 10px 0; border-radius: 5px; }}
                    table {{ width: 100%; border-collapse: collapse; }}
                    td, th {{ border: 1px solid #2a4a3a; padding: 8px; text-align: left; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">
                        <h1>File Analysis Report</h1>
                        <h2 class="risk-{analysis_result['risk_level'].lower()}">Risk Level: {analysis_result['risk_level']}</h2>
                    </div>
                    
                    <div class="section">
                        <h3>File Information</h3>
                        <table>
                            <tr><td>Filename</td><td>{analysis_result['filename']}</td></tr>
                            <tr><td>Timestamp</td><td>{analysis_result['timestamp']}</td></tr>
                            <tr><td>Risk Score</td><td>{analysis_result['risk_score']}/100</td></tr>
                            <tr><td>File Size</td><td>{analysis_result['file_size']} bytes</td></tr>
                            <tr><td>Entropy</td><td>{analysis_result['entropy']:.2f}</td></tr>
                        </table>
                    </div>
                    
                    <div class="section">
                        <h3>File Hashes</h3>
                        <table>
                            <tr><td>MD5</td><td>{analysis_result['hash']['md5']}</td></tr>
                            <tr><td>SHA1</td><td>{analysis_result['hash']['sha1']}</td></tr>
                            <tr><td>SHA256</td><td>{analysis_result['hash']['sha256']}</td></tr>
                        </table>
                    </div>
                    
                    <div class="section">
                        <h3>Suspicious Patterns ({len(analysis_result.get('suspicious_patterns', []))})</h3>
                        <table>
                            <tr><th>Type</th><th>Pattern</th><th>Severity</th></tr>
            """
            for p in analysis_result.get('suspicious_patterns', []):
                html_data += f"<tr><td>{p['type']}</td><td>{p['pattern']}</td><td>{p['severity']}</td></tr>"
            
            html_data += """
                        </table>
                    </div>
                </div>
            </body>
            </html>
            """
            b64_html = base64.b64encode(html_data.encode()).decode()
            href_html = f'<a href="data:text/html;base64,{b64_html}" download="{analysis_result["filename"]}_report.html" class="download-btn">📥 HTML</a>'
            st.markdown(href_html, unsafe_allow_html=True)
        
        # Suspicious patterns found
        if analysis_result['suspicious_patterns']:
            st.markdown("### 🔍 Suspicious Patterns Detected")
            patterns_df = pd.DataFrame(analysis_result['suspicious_patterns'])
            st.dataframe(patterns_df, use_container_width=True)
        
        # AI Analysis
        if 'ai_analysis' in analysis_result:
            ai = analysis_result['ai_analysis']
            
            # Threat Level
            threat = analysis_result['risk_level']
            if threat == 'CRITICAL':
                st.error(f"### ⚠️ THREAT LEVEL: {threat}")
            elif threat == 'HIGH':
                st.warning(f"### ⚠️ THREAT LEVEL: {threat}")
            elif threat == 'MEDIUM':
                st.info(f"### THREAT LEVEL: {threat}")
            else:
                st.success(f"### THREAT LEVEL: {threat}")
            
            # Attack Type
            st.success(f"### 🎯 Attack Type: {ai.get('attack_type', 'Unknown')}")
            
            # Context
            st.markdown("### 📋 Context")
            context = ai.get('context', 'No context available')
            if isinstance(context, str):
                sentences = context.replace('. ', '.\n').split('\n')
                for sentence in sentences[:6]:
                    if sentence.strip():
                        clean = sentence.strip()
                        if clean and clean[0].isdigit() and '.' in clean[:3]:
                            clean = clean[clean.find('.')+1:].strip()
                        st.markdown(f"• {clean}")
            
            # Prevention
            st.markdown("### 🛡️ Prevention")
            prevention = ai.get('prevention', 'No prevention steps available')
            if isinstance(prevention, str):
                steps = prevention.replace('. ', '.\n').split('\n')
                for i, step in enumerate(steps[:6], 1):
                    if step.strip():
                        clean = step.strip()
                        if clean and clean[0].isdigit() and '.' in clean[:3]:
                            clean = clean[clean.find('.')+1:].strip()
                        st.markdown(f"**{i}.** {clean}")

# ===== TAB 3: Reports =====
with tab3:
    st.markdown("## 📊 Analysis Reports")
    
    if st.session_state.file_analysis_results:
        # Create a comprehensive report dataframe
        report_list = []
        for result in st.session_state.file_analysis_results:
            report_list.append({
                'Timestamp': result['timestamp'],
                'Filename': result['filename'],
                'Risk Level': result['risk_level'],
                'Risk Score': result['risk_score'],
                'File Size (bytes)': result['file_size'],
                'Entropy': f"{result['entropy']:.2f}",
                'Suspicious Patterns': len(result.get('suspicious_patterns', [])),
                'Attack Summary': result.get('attack_summary', 'None'),
                'MD5': result['hash']['md5'],
                'SHA256': result['hash']['sha256'][:16] + '...'
            })
        
        report_df = pd.DataFrame(report_list)
        
        # Display the report
        st.dataframe(report_df, use_container_width=True)
        
        st.markdown('<div class="section-header">📥 Download All Reports</div>', unsafe_allow_html=True)
        
        # Download buttons
        st.markdown('<div class="download-container">', unsafe_allow_html=True)
        
        # Function to generate PDF report
        def generate_pdf_report():
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
            elements = []
            
            # Title
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.green,
                alignment=1,
                spaceAfter=30
            )
            elements.append(Paragraph("AI Intrusion Detection System - Analysis Report", title_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Date
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=10,
                alignment=1,
                spaceAfter=20
            )
            elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", date_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Summary stats
            stats_data = [
                ['Total Files', 'Critical', 'High', 'Medium', 'Low'],
                [
                    str(len(report_list)),
                    str(sum(1 for r in report_list if r['Risk Level'] == 'CRITICAL')),
                    str(sum(1 for r in report_list if r['Risk Level'] == 'HIGH')),
                    str(sum(1 for r in report_list if r['Risk Level'] == 'MEDIUM')),
                    str(sum(1 for r in report_list if r['Risk Level'] == 'LOW'))
                ]
            ]
            stats_table = Table(stats_data, colWidths=[1.2*inch]*5)
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(stats_table)
            elements.append(Spacer(1, 0.3*inch))
            
            # Detailed results table
            table_data = [['Timestamp', 'Filename', 'Risk Level', 'Score', 'Patterns']]
            for r in report_list[-20:]:  # Last 20 entries
                table_data.append([
                    r['Timestamp'][11:19],
                    r['Filename'][:20] + '...' if len(r['Filename']) > 20 else r['Filename'],
                    r['Risk Level'],
                    str(r['Risk Score']),
                    str(r['Suspicious Patterns'])
                ])
            
            detailed_table = Table(table_data, colWidths=[1*inch, 2*inch, 0.8*inch, 0.6*inch, 0.8*inch])
            detailed_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(detailed_table)
            
            doc.build(elements)
            buffer.seek(0)
            return buffer.getvalue()
        
        # Function to generate DOCX report
        def generate_docx_report():
            doc = Document()
            
            # Title
            title = doc.add_heading('AI Intrusion Detection System - Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_heading('Summary Statistics', level=1)
            
            # Summary table
            summary_table = doc.add_table(rows=2, cols=5)
            summary_table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = ['Total Files', 'Critical', 'High', 'Medium', 'Low']
            for i, header in enumerate(headers):
                summary_table.cell(0, i).text = header
            
            # Data
            data = [
                str(len(report_list)),
                str(sum(1 for r in report_list if r['Risk Level'] == 'CRITICAL')),
                str(sum(1 for r in report_list if r['Risk Level'] == 'HIGH')),
                str(sum(1 for r in report_list if r['Risk Level'] == 'MEDIUM')),
                str(sum(1 for r in report_list if r['Risk Level'] == 'LOW'))
            ]
            for i, value in enumerate(data):
                summary_table.cell(1, i).text = value
            
            doc.add_heading('Detailed Analysis Results', level=1)
            
            # Detailed table
            detailed_table = doc.add_table(rows=1, cols=5)
            detailed_table.style = 'Light Grid Accent 1'
            
            # Headers
            header_cells = detailed_table.rows[0].cells
            header_cells[0].text = 'Time'
            header_cells[1].text = 'Filename'
            header_cells[2].text = 'Risk Level'
            header_cells[3].text = 'Score'
            header_cells[4].text = 'Patterns'
            
            # Data
            for r in report_list[-20:]:
                row_cells = detailed_table.add_row().cells
                row_cells[0].text = r['Timestamp'][11:19]
                row_cells[1].text = r['Filename'][:30] + '...' if len(r['Filename']) > 30 else r['Filename']
                row_cells[2].text = r['Risk Level']
                row_cells[3].text = str(r['Risk Score'])
                row_cells[4].text = str(r['Suspicious Patterns'])
            
            # Save to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            return docx_buffer.getvalue()
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # CSV Download
            csv_data = report_df.to_csv(index=False)
            b64 = base64.b64encode(csv_data.encode()).decode()
            href = f'<a href="data:file/csv;base64,{b64}" download="analysis_report.csv" class="download-btn">📥 Download CSV</a>'
            st.markdown(href, unsafe_allow_html=True)
        
        with col2:
            # PDF Download
            try:
                pdf_data = generate_pdf_report()
                b64_pdf = base64.b64encode(pdf_data).decode()
                href_pdf = f'<a href="data:application/pdf;base64,{b64_pdf}" download="analysis_report.pdf" class="download-btn">📥 Download PDF</a>'
                st.markdown(href_pdf, unsafe_allow_html=True)
            except Exception as e:
                st.error(f"PDF generation failed: {e}")
        
        with col3:
            # DOCX Download
            try:
                docx_data = generate_docx_report()
                b64_docx = base64.b64encode(docx_data).decode()
                href_docx = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="analysis_report.docx" class="download-btn">📥 Download DOCX</a>'
                st.markdown(href_docx, unsafe_allow_html=True)
            except Exception as e:
                st.error(f"DOCX generation failed: {e}")
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Statistics
        st.markdown("### 📈 Statistics")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Files", len(report_list))
        with col2:
            critical = sum(1 for r in report_list if r['Risk Level'] == 'CRITICAL')
            st.metric("Critical", critical)
        with col3:
            high = sum(1
